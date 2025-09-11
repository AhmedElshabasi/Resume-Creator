from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()

# Add bold "Experience" text
p = doc.add_paragraph()
run = p.add_run("Experience")
run.bold = True
run.font.size = Pt(14)

# Add bottom border to simulate the horizontal line
p_par = p._element
p_borders = OxmlElement('w:pBdr')
bottom_border = OxmlElement('w:bottom')
bottom_border.set(qn('w:val'), 'single')
bottom_border.set(qn('w:sz'), '6')  # thickness
bottom_border.set(qn('w:space'), '1')
bottom_border.set(qn('w:color'), 'auto')
p_borders.append(bottom_border)
p_par.get_or_add_pPr().append(p_borders)

# Save the document
doc.save("styled_header_line.docx")
