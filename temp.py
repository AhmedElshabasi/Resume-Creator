import os
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
import re
from datetime import datetime

# Load .env file
load_dotenv()

client = OpenAI(
    api_key = os.getenv("OPENAI_API_KEY"),
)

with open("input.txt", "r", encoding="utf-8") as file:
    job_descrip = file.read()

completion = client.chat.completions.create(
    model="gpt-4o",
    messages=[
        {
            "role": "user",
            "content": f"""Ahmed Elshabasi
(587)-435-9647 – Calgary, AB – ahmed.elshabasi@ucalgary.ca – LinkedIn: linkedin.com/in/ahmed-elshabasi/
Skills
•	Programming languages: Java, Python, JavaScript, C, C++, HTML, CSS, React, Node.JS, Express.JS, SQL
•	Software tools: VS Code, Git, Github, Gitlab, Unity, Unreal Engine
•	Algorithm and Data Structures: Studied different algorithms and structures in university
•	Professional Skills: Adaptability, Communication, Detail-oriented, Leadership, and Time Management
Experience
Undergraduate Research Assistant (Node, React, JS)					                           May 2024 – Sep 2024
University of Calgary,     Calgary, AB
•	Developed an automated workflow using Node and React for extracting in-depth information about the data in the corpus completing tasks within set deadlines.
•	Collected first-person videos, spoken recordings, and biometric data for a corpus of information needs for outdoor running.
•	Quickly learned new tools and technologies (Node, React, and smart capture devices) to automate data analysis workflows.
Projects
Self-Checkout Machine Software (Java)								          Sep 2023 – Dec 2023
•	Collaborated with a team of 20 to designed and develop the software for a self-checkout machine
•	Focused on user-friendly interface design and efficient transaction handling to ensure smooth customer experience.
•	Integrated functionalities that simulate real-world use cases.
Educational Assessment Web App (JS, CSS, HTML)						            Jan 2024 – Apr 2024
•	Collaborated with a team of 5 to design a web application for dynamic assessments.
•	Implemented functionality to generate random questions per session, offering immediate grading and feedback.
•	Emphasized designing a user-friendly interface to provide smooth navigation and create an engaging test experience.
Full-stack Financial Assistant| Hackathon Project (Node, React, JS)	   		     	                              Feb 2024
•	Led a team of 4 to build a full-stack prompt-based financial assistant.
•	 Used ChatGPT’s API for real-time financial insights and assistance.
•	Ensured seamless deployment within the time constraints of the hackathon (24 hours).
Education
Bachelor of Computer Science									   2022 – 2026 [Expected]
University of Calgary,  Calgary, AB									    GPA: 3.68
•	Awards:
o	PURE award
o	President’s Admission Scholarship
o	University of Calgary International Undergraduate Award
•	Certificates:
o	Google Cybersecurity Professional Certificate		
o	Ready for Research Micro credential
Volunteering
Setup Crew											          Jan 2022 – May 2022
G.N.P. Hospital,    Jeddah, Saudi Arabia
•	Assisted medical students with a staff team of 5 and documented the students’ progress.
•	Collaboratively smoothed the experience for the students by getting their feedback and answering inquiries.
•	Recorded students’ progress and managed their attendance for their hospital’s academy training.
Executive Team Member									          Dec 2021 - Apr 2022
Model United Nations (MUN) at Dar Jana International School
•	Managed and participated with the MUN executive team in the organization of documents and preparation of participants
•	Prepared the hall in which the event takes place and ensured that the event proceeded as expected.
•	As one of the spokespersons (chairmen) of the event, fulfilled my duties of planning and managing the procedures of the delegates.
Extracurricular Clubs
Event Coordinator										         Sep 2023 – May 2024
Infosec Club											
•	Coordinated hands-on security labs and ethical hacking workshops, providing practical learning experiences for members.
Member											         Sep 2023 – May 2024
Competitive Programming Club									 
•	Engaged in practical programming challenges, honing my problem-solving skills through regular workshop attendance.




this is my resume and bs coverletter keep that in mind cause i am about to ask you to do stuff

Here is the job description: {job_descrip}

i want you to infiltrate my experiences with details that aid my core skills and support my ability to do the major aspects of what's required from the job description (it doesn't have to be legitimate but it has to be powerful), then create a resume.

Write me a resume i want it in a very strict format because the answer would be used to fill a word document template, so it has to be in a specific format. The format is as follows:

Skills
• Programming languages: {{programming_languages}}
• Software tools: {{software_tools}}
• Algorithm and Data Structures: {{algorithms_and_structures}}
• Professional Skills: {{professional_skills}}

Experience
{{position_1}} ({{skills_1}})					                           
{{organization_1}}, {{location_1}}					                           
{{start_date_1}} – {{end_date_1}}
• {{bullet_1_1}}
• {{bullet_1_2}}
• {{bullet_1_3}}
• {{bullet_1_4}}

{{position_2}} ({{program_type_2}})
{{organization_2}}, {{location_2}}
{{start_date_2}} – {{end_date_2}}
• {{bullet_2_1}}
• {{bullet_2_2}}
• {{bullet_2_3}}
• {{bullet_2_4}}
• {{bullet_2_5}}
• {{bullet_2_6}}

Projects
{{project_1_title}} ({{project_1_tech}})
{{project_1_start}} – {{project_1_end}}
• {{project_1_bullet_1}}
• {{project_1_bullet_2}}
• {{project_1_bullet_3}}

{{project_2_title}} ({{project_2_tech}})
{{project_2_start}} – {{project_2_end}}
• {{project_2_bullet_1}}
• {{project_2_bullet_2}}
• {{project_2_bullet_3}}

{{project_3_title}} | Hackathon Project ({{project_3_tech}})
{{project_3_date}}
• {{project_3_bullet_1}}
• {{project_3_bullet_2}}
• {{project_3_bullet_3}}

{{project_4_title}} ({{project_4_tech}})
{{project_4_start}} – {{project_4_end}}
• {{project_4_bullet_1}}
• {{project_4_bullet_2}}
• {{project_4_bullet_3}}
• {{project_4_bullet_4}}
• {{project_4_bullet_5}}
• {{project_4_bullet_6}}
• {{project_4_bullet_7}}
• {{project_4_bullet_8}}

{{project_5_title}}
Team Size: {{project_5_team_size}} | Role: {{project_5_role}}
• {{project_5_bullet_1}}
• {{project_5_bullet_2}}
• {{project_5_bullet_3}}
• {{project_5_bullet_4}}
• {{project_5_bullet_5}}

Education
{{degree}}									   
{{education_institution}}, {{education_location}}					   
{{education_duration}} | GPA: {{gpa}}
• Awards:
o {{award_1}}
o {{award_2}}
o {{award_3}}
• Certificates:
o {{cert_1}}
o {{cert_2}}
o {{cert_3}}
o {{cert_4}}

Volunteering
{{volunteer_role_1}}
{{volunteer_org_1}}, {{volunteer_loc_1}}					   
{{volunteer_date_1}}
• {{volunteer_1_bullet_1}}
• {{volunteer_1_bullet_2}}
• {{volunteer_1_bullet_3}}

{{volunteer_role_2}}
{{volunteer_org_2}}					   
{{volunteer_date_2}}
• {{volunteer_2_bullet_1}}
• {{volunteer_2_bullet_2}}

Extracurricular Clubs
{{club_role_1}}					   
{{club_1_name}}					   
{{club_1_date}}
• {{club_1_bullet_1}}

{{club_role_2}}					   
{{club_2_name}}					   
{{club_2_date}}
• {{club_2_bullet_1}}

So, the format of your answer should look something like this:
{{programming_languages}}:

"""
        }
    ]
)

api_response = completion.choices[0].message.content

# Extract values using regex dynamically
data = {}
matches = re.findall(r"\[(.*?)\]:\[\[\[(.*?)\]\]\]", api_response, re.DOTALL)  # re.DOTALL captures multi-line content
for key, value in matches:
    data[f"[{key}]"] = value.strip()  # Strip to remove unwanted spaces/newlines

# Add today's date
today_str = datetime.now().strftime("%d/%m/%Y")
data["[Date123]"] = today_str

print(f"data: {data}")

# Load the existing Word document (template)
doc = Document("coverletter.docx")  # Make sure your template exists

# Replace placeholders in normal text (single-line values)
for para in doc.paragraphs:
    for key, value in data.items():
        if key in para.text:
            if key == "[Content123]":
                # Special handling for multi-line content
                para.clear()  # Remove placeholder text
                for line in value.split("\n"):  # Keep each paragraph separate
                    para.add_run(line)
                    para.add_run("\n")  # Add a newline for spacing
            else:
                para.text = para.text.replace(key, value)

# Save the updated document
doc.save("AhmedElshabasi_CoverLetter.docx")

print("Document updated successfully!")
