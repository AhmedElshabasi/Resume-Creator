import os
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
import re
from datetime import datetime
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
import json
import json
import re
from docx.shared import Inches, Pt
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

# --- tiny helpers  ---
def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)                 
    p = Paragraph(new_p, paragraph._parent)
    if style:
        p.style = style
    if text:
        p.add_run(text)
    return p

def bulletify(par: Paragraph, text: str):
    par.paragraph_format.left_indent = Inches(0.25)
    par.paragraph_format.first_line_indent = Inches(-0.15)
    par.add_run("• " + text)
    apply_spacing_and_size(par)     


def first_existing_style(doc: Document, candidates: list[str]):
    """Return the first style object that exists in doc.styles from candidates, else None."""
    for name in candidates:
        try:
            return doc.styles[name]
        except KeyError:
            continue
    return None

def apply_spacing_and_size(p: Paragraph, font_pt: int = 10):
    pf = p.paragraph_format
    pf.line_spacing = 1            # single line
    pf.space_before = Pt(3)        # 3 pt before
    pf.space_after = Pt(0)         # 0 pt after
    # Optional (recommended to keep layout tight):
    # pf.space_after = Pt(0)

    for r in p.runs:
        r.font.size = Pt(font_pt)







def replace_placeholder_with_bullets(doc: Document, placeholder: str, items: list[str]) -> bool:
    """
    Replace a placeholder paragraph (e.g. "[skills123]") with a bullet list.
    Uses an existing bullet/list style if found; otherwise falls back to a
    manual bullet glyph and hanging indent.
    """
    bullet_style_obj = first_existing_style(
        doc,
        # Try common English names first; add your own if needed
        ["List Bullet", "List Paragraph", "Bulleted List", "Bullet", "Normal"]
    )

    def iter_paragraphs(d: Document):
        for p in d.paragraphs:
            yield p
        for tbl in d.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    for p in iter_paragraphs(doc):
        if placeholder in p.text:
            if not items:
                # nothing to insert, just clear placeholder
                set_paragraph_text(p, "")
                return True

            # Decide rendering path: real style vs manual bullets
            use_manual_bullets = (bullet_style_obj is None or bullet_style_obj.name in ("Normal", "List Paragraph"))

            def style_as_bullet(par):
                if bullet_style_obj is not None:
                    par.style = bullet_style_obj
                if use_manual_bullets:
                    # manual bullet look
                    par.paragraph_format.left_indent = Inches(0.25)
                    par.paragraph_format.first_line_indent = Inches(-0.15)

            # First item replaces the placeholder line
            style_as_bullet(p)
            set_paragraph_text(p, (("• " if use_manual_bullets else "") + items[0]))
            apply_spacing_and_size(p)

            # Remaining items
            anchor = p
            for it in items[1:]:
                anchor = insert_paragraph_after(anchor)
                style_as_bullet(anchor)
                set_paragraph_text(anchor, (("• " if use_manual_bullets else "") + it))
                apply_spacing_and_size(anchor)
            return True

    return False

# --- main function ---
def replace_experience_placeholder(doc, placeholder: str, items: list[dict]) -> bool:
    # find the placeholder paragraph
    target_p = None
    for p in doc.paragraphs:
        if placeholder in p.text:
            target_p = p
            break
    if target_p is None:
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if placeholder in p.text:
                            target_p = p
                            break
                if target_p: break
            if target_p: break
    if target_p is None:
        return False

    # clear placeholder to reuse the same paragraph for the first header
    for r in list(target_p.runs):
        r._element.getparent().remove(r._element)

    anchor = target_p
    first = True

    for entry in items:
        role   = entry.get("role", "").strip()
        dates  = entry.get("dates", "").strip()
        loc    = entry.get("location", "").strip()
        detail = entry.get("details", []) or []

        # HEADER line (role left, dates right)
        header_p = anchor if first else insert_paragraph_after(anchor)
        first = False

        pf = header_p.paragraph_format
        pf.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)

        run_role = header_p.add_run("❖ " + role)
        run_role.bold = True
        header_p.add_run("\t" + dates)

        apply_spacing_and_size(header_p)        # <-- make header 10pt, single, 3pt before

        # LOCATION line
        loc_p = insert_paragraph_after(header_p, loc)
        for r in loc_p.runs:
            r.italic = True
        apply_spacing_and_size(loc_p)           # <-- format location line

        # DETAILS bullets (indented)
        prev = loc_p
        for d in detail:
            bp = insert_paragraph_after(prev)
            bulletify(bp, d)                    # bulletify already applies spacing/size
            prev = bp

        anchor = prev

    return True

def replace_projects_placeholder(doc: Document, placeholder: str, items: list[dict]) -> bool:
    # find placeholder paragraph (search body and tables)
    target_p = None
    for p in doc.paragraphs:
        if placeholder in p.text:
            target_p = p
            break
    if target_p is None:
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if placeholder in p.text:
                            target_p = p
                            break
                if target_p: break
            if target_p: break
    if target_p is None:
        return False

    # wipe placeholder
    for r in list(target_p.runs):
        r._element.getparent().remove(r._element)

    anchor = target_p
    first = True

    for proj in items:
        title   = proj.get("title", "").strip()
        dates   = proj.get("dates", "").strip()
        stack   = proj.get("stack", "").strip()
        details = proj.get("details", []) or []
        link    = proj.get("link", "").strip()

        # HEADER: title left (bold + diamond), dates right (tab stop)
        header_p = anchor if first else insert_paragraph_after(anchor)
        first = False

        pf = header_p.paragraph_format
        pf.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)

        run_title = header_p.add_run("❖ " + title)
        run_title.bold = True
        if dates:
            header_p.add_run("\t" + dates)

        apply_spacing_and_size(header_p)  # <-- make header 10pt, single, 3pt before

        # SUBTITLE: stack (+ optional link) in italics
        subtitle_parts = []
        if stack:
            subtitle_parts.append(stack)
        if link:
            subtitle_parts.append(link)
        if subtitle_parts:
            sub_p = insert_paragraph_after(header_p, " – ".join(subtitle_parts))
            for r in sub_p.runs:
                r.italic = True
            apply_spacing_and_size(sub_p)   # <-- subtitle formatting
        else:
            sub_p = header_p  # no subtitle; bullets go directly after header

        # DETAILS bullets (bulletify already applies spacing/size)
        prev = sub_p
        for d in details:
            bp = insert_paragraph_after(prev)
            bulletify(bp, d)                 # applies 10pt, single, 3pt before
            prev = bp

        anchor = prev  # next project continues after last bullet

    return True

def replace_education_placeholder(doc: Document, placeholder: str, items: list[dict]) -> bool:
    # find placeholder
    target_p = None
    for p in doc.paragraphs:
        if placeholder in p.text:
            target_p = p
            break
    if target_p is None:
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if placeholder in p.text:
                            target_p = p
                            break
                if target_p: break
            if target_p: break
    if target_p is None:
        return False

    # clear placeholder
    for r in list(target_p.runs):
        r._element.getparent().remove(r._element)

    anchor = target_p
    first = True

    for edu in items:
        degree  = edu.get("degree", "").strip()
        dates   = edu.get("dates", "").strip()
        loc     = edu.get("location", "").strip()
        details = edu.get("details", []) or []

        # HEADER: degree left (bold), dates right-aligned
        header_p = anchor if first else insert_paragraph_after(anchor)
        first = False

        pf = header_p.paragraph_format
        pf.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)

        run_degree = header_p.add_run(degree)
        run_degree.bold = True
        if dates:
            header_p.add_run("\t" + dates)

        apply_spacing_and_size(header_p)            # ← added

        # LOCATION line
        loc_p = insert_paragraph_after(header_p, loc)
        for r in loc_p.runs:
            r.italic = True
        apply_spacing_and_size(loc_p)               # ← added

        # DETAILS bullets (bulletify already applies spacing/size)
        prev = loc_p
        for d in details:
            bp = insert_paragraph_after(prev)
            bulletify(bp, d)                        # bulletify calls apply_spacing_and_size
            prev = bp

        anchor = prev

    return True


def extract_json(text: str) -> str:
    """
    Try a few strategies to pull a valid JSON object from a model response.
    1) If it's inside a code fence ```json ... ```, grab that.
    2) If it looks like `resume_data = { ... }`, grab the {...}.
    3) Fallback: take the first balanced {...} block.
    """
    if not text or not text.strip():
        raise ValueError("Empty API response")

    # 1) ```json ... ``` or ``` ... ```
    m = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, flags=re.DOTALL)
    if m:
        return m.group(1)

    # 2) resume_data = { ... }
    m = re.search(r"resume_data\s*=\s*(\{.*\})", text, flags=re.DOTALL)
    if m:
        return m.group(1)

    # 3) first balanced { ... } block (greedy from first { to last })
    first = text.find("{")
    last = text.rfind("}")
    if first != -1 and last != -1 and last > first:
        return text[first:last+1]

    # If all else fails
    raise ValueError("Could not find JSON object in response")


# Load .env file
load_dotenv()

client = OpenAI(
    api_key = os.getenv("OPENAI_API_KEY"),
)

with open("input.txt", "r", encoding="utf-8") as file:
    job_descrip = file.read()

print("Job description loaded. and about to call the api")

completion = client.chat.completions.create(
    model="gpt-4o",
    messages=[
        {
            "role": "user",
            "content": f"""Ahmed Elshabasi
(587)-435-9647 – Calgary, AB – ahmed.elshabasi@ucalgary.ca – LinkedIn: linkedin.com/in/ahmed-elshabasi/
Skills
•	Programming languages: Java, Python, JavaScript, C, C++, HTML, CSS, React, Node.JS, Express.JS, SQL
•	Software tools: VS Code, Git, Github, Gitlab, Docker, Unity, Unreal Engine, MS Office
•	Algorithm and Data Structures: Studied different algorithms and structures in university
•	Professional Skills: Adaptability, Communication, Detail-oriented, Leadership, and Time Management
Experience
Undergraduate Research Assistant (Node, React, JS)					                           May 2024 – Sep 2024
University of Calgary,     Calgary, AB
•	Developed an automated workflow using Node and React for extracting in-depth information about the data in the corpus and creating data visualizations, completing tasks within set deadlines.
•	Developed an automated transcription tool using OpenAI’s API Whisper that transcribed hundreds of videos to allow for data analysis.
•	Collected first-person videos, spoken recordings, and biometric data for a corpus of information needs for outdoor running.
•	Quickly learned new tools and technologies (Node, React, and smart capture devices) to automate data analysis workflows.
Summer Intern – Sharpen Up Internship Program (Rotational)					       Jun 2025 – Aug 2025
Viatris Egypt, Cairo, Egypt
•	Rotated across departments including Regulatory Affairs, Quality, Medical, and Supply Chain to gain a broad understanding of pharmaceutical operations.
•	Collaborated with the Data & Analytics team to assist in developing dashboards for supply chain performance monitoring using Excel and basic SQL queries.
•	Supported the Regulatory Affairs team by preparing documentation for product registration and compliance tracking.
•	Participated in pharmacovigilance training, learning to evaluate case safety reports and maintain safety databases.
•	Delivered a final group presentation analyzing business process improvements using data collected during the internship.
•	Attended workshops on pharmaceutical ethics, global compliance, and the drug approval lifecycle.

Projects
Self-Checkout Machine Software (Java)								          Sep 2023 – Dec 2023
•	Collaborated with a team of 20 to designed and develop the software for a self-checkout machine
•	Focused on user-friendly interface design and efficient transaction handling to ensure smooth customer experience.
•	Integrated functionalities that simulate real-world use cases.
Educational Assessment Web App (JS, CSS, HTML)						            Jan 2024 – Apr 2024
•	Collaborated with a team of 5 to design a web application for dynamic assessments.
•	Implemented functionality to generate random questions per session, offering immediate grading and feedback.
•	Emphasized designing a user-friendly interface to provide smooth navigation and create an engaging test experience.
Full-stack Financial Assistant| Hackathon Project (Node, React, JS)	   		     	                              Feb 2024
•	Led a team of 4 to build a full-stack prompt-based financial assistant.
•	Used ChatGPT’s API for real-time financial insights and assistance.
•	Ensured seamless deployment within the time constraints of the hackathon (24 hours).
SceneBook: Multi-Theatre Aggregation & Ticketing Platform (PostgreSQL, Node.js, React)	Jan 2025 – Apr 2025

•	Built a full-stack aggregation platform for movie ticketing that unified listings, showtimes, and booking flows across multiple theatre companies (e.g., Cineplex, Landmark).
•	Integrated dynamic user and admin experiences, including account management, ticket booking, payment (Card/PayPal), and seat reservation.
•	Designed and implemented a normalized PostgreSQL database schema supporting multi-tenant theatres, admin roles, and a relational audit trail.
•	Developed SQL-driven APIs for complex operations like rating moderation, payment verification, and movie request handling, using node-postgres.
•	Created robust admin dashboards for movie listing management, theatre configuration, and comment moderation based on user roles.
•	Engineered frontend features with React, such as location-based search, company filters, and booking flows with seat previews and payment validation.
•	Documented architecture using HIPO diagrams, DFDs, and ERDs, and published a complete user guide for navigation and deployment.
•	GitHub: github.com/AhmedElshabasi/CPSC471W25Project
EventEcho – Full-Stack Event Management Web Application					Sep 2024 – Dec 2024
	Team Size: 5 developers | Role: Backend & Systems Integration Developer
•	Designed and implemented RESTful APIs using Node.js and Express to support event creation, registration, friend management, and admin moderation features.
•	Integrated PostgreSQL as the primary database, structured normalized schemas, and managed migration/seeding scripts for scalable data persistence.
•	Implemented token-based authentication with JWT to handle user login, session security, and admin privilege access (including user bans and event removals).
•	Developed dynamic backend logic to support different user roles (guest, user, admin) and features like private/public events, event invites, and Cloudinary-hosted image uploads.
•	Collaborated with frontend developers using Git/GitHub and Docker Compose to ensure full-stack alignment and CI-ready deployment in a mobile-first web architecture (React + MUI).

Education
Bachelor of Computer Science									   2022 – 2026 [Expected]
University of Calgary,  Calgary, AB									    GPA: 3.68
•	Awards:
o	PURE (Program for Undergraduate Research Experience) award
o	President’s Admission Scholarship
o	University of Calgary International Undergraduate Award
•	Certificates:
o	CompTIA Security+
o	Pursuing CompTIA Network+
o	Google Cybersecurity Professional Certificate		
o	Ready for Research Micro credential
Volunteering
Setup Crew											          Jan 2022 – May 2022
G.N.P. Hospital,    Jeddah, Saudi Arabia
•	Assisted medical students with a staff team of 5 and documented the students’ progress.
•	Collaboratively smoothed the experience for the students by getting their feedback and answering inquiries.
•	Recorded students’ progress and managed their attendance for their hospital’s academy training.
Executive Team Member									          Dec 2021 - Apr 2022
Model United Nations (MUN) at Dar Jana International School
•	Prepared the hall in which the event takes place and ensured that the event proceeded as expected.
•	As one of the spokespersons (chairmen) of the event, fulfilled my duties of planning and managing the procedures of the delegates.
Extracurricular Clubs
Event Coordinator										         Sep 2023 – May 2024
Infosec Club											
•	Coordinated hands-on security labs and ethical hacking workshops, providing practical learning experiences for members.
Member											         Sep 2023 – May 2024
Competitive Programming Club									 
•	Engaged in practical programming challenges, honing my problem-solving skills through regular workshop attendance.




this is my resume and bs coverletter keep that in mind cause i am about to ask you to do stuff

Here is the job description: {job_descrip}

i want you to infiltrate my experiences with details that aid my core skills and support my ability to do the major aspects of what's required from the job description (it doesn't have to be legitimate but it has to be powerful), then create a resume.

i want you to output your answer in a very strict format of a json array that looks like:

resume_data = {{
    "skills123": [
        "result1",
        "result2",
        "result3",
        "result4.... and so on as necessary"
    ],
     "education123": [
    {{
      "degree": "Bachelor of Computer Science",
      "dates": "Sep 2022 – Apr 2026 [Expected]",
      "location": "University of Calgary – Calgary, Alberta",
      "details (extra things like certifications, awards, relevant courses)": [
      e.g.:
        "Certifications: CompTIA Sec+ | Pursuing CompTIA Network+ | Google Cybersecurity Professional Certificate",
        "Awards: PURE (Program for Undergraduate Research Experience) award, University of Calgary International Undergraduate Award",
        "Relevant Courses: Operating Systems, Networking, Data Structures, Cybersecurity, Software Engineering"}}
      ],
    "experience123": [
        {{
        "role": "role1",
        "dates": "dates1",
        "location": "location1"
        "details": [
            "result1",
            "result2",
            "result3",
            "result4.... and so on as necessary"}}
        ],
    e.g. for projects:
    "projects123": [
    {{
      "title": "	EventEcho – Full-stack Role-Based Web Platform ",
      "dates": "Sep 2024 – Dec 2024",
      "stack": "Node.js, React, JWT, PostgreSQL, Docker",
      "details": [
        "Implemented token-based authentication with role-based access for admins, guests, and users.",
        "Designed and maintained secure user registration and login APIs aligned with least-privilege access principles.",
        "Documented access workflows and helped define edge-case handling for login failures and unauthorized actions."
      ],
    }},
}}
"""
        }
    ]
)

api_response = completion.choices[0].message.content
print("this is the api response:", api_response)

# ---------- helpers ----------

def set_paragraph_text(p: Paragraph, text: str):
    """Replace all runs in a paragraph with a single run containing `text`."""
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(text)


# ---------- usage ----------


api_response = completion.choices[0].message.content
json_str = extract_json(api_response)
resume_data = json.loads(json_str)

doc = Document("resume_template.docx")           # your .docx version of the 2nd screenshot
replace_placeholder_with_bullets(doc, "[skills123]", resume_data["skills123"])
doc.save("resume_filled_skills.docx")

doc = Document("resume_filled_skills.docx")
replace_experience_placeholder(doc, "[experience123]", resume_data["experience123"])
doc.save("resume_filled_experience.docx")

doc = Document("resume_filled_experience.docx")  # or resume_template.docx if you’re doing projects first
replace_projects_placeholder(doc, "[projects123]", resume_data["projects123"])
doc.save("resume_filled_projects.docx")

doc = Document("resume_filled_projects.docx")  # or resume_template if fresh
replace_education_placeholder(doc, "[education123]", resume_data["education123"])
doc.save("resume_filled_education.docx")

import os

# open the file automatically (Windows only)
os.startfile("resume_filled_education.docx")

